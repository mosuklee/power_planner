# ------------------------------------------------------
# ---------------------- main.py -----------------------
# 작업환경 : 
# ------------------------------------------------------
from time import time
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi
from PyQt5 import QtCore,QtWidgets
from PyQt5.QtCore import Qt
from PyQt5.QtCore import QSize
from PyQt5.QtWidgets import QFileDialog, QLabel
from PyQt5.QtWidgets import QAction, QMainWindow, QMessageBox
from matplotlib.backends.backend_qt5agg import (NavigationToolbar2QT as NavigationToolbar)
from matplotlib import font_manager, rc
# url 연결결
from PyQt5.QtCore import QUrl
from PyQt5.QtGui import QDesktopServices, QFont  # QFont 추가

# 4. Report를 만들기 위한 모듈
from docx2pdf import convert

import os
import sys
import math

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt  # 그래프 시각화 패키지

from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException, TimeoutException
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.action_chains import ActionChains 
from selenium.common.exceptions import ElementClickInterceptedException, TimeoutException 


from selenium.webdriver.support.ui import WebDriverWait 
from selenium.webdriver.support import expected_conditions as EC 
from selenium.webdriver.common.by import By

from selenium.common.exceptions import TimeoutException, ElementClickInterceptedException


from bs4 import BeautifulSoup
from urllib.request import urlopen
import urllib.parse
import pandas as pd
import re
import time
from selenium.webdriver.common.by import By

from PyQt5.QtWidgets import QHeaderView

# DPI 스케일링 설정
def main():
    # QApplication 객체 생성 전에 속성을 설정합니다.
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)

    app = QApplication(sys.argv)

    # DPI를 고려한 기본 폰트 크기 설정
    dpi = app.primaryScreen().logicalDotsPerInch()
    font_size = int(10 * dpi / 96)  # 96은 기본 DPI
    app.setFont(QFont("Arial", font_size))
    
    window = MatplotlibWidget()
    window.show()
    sys.exit(app.exec_())


class MatplotlibWidget(QMainWindow):

    def __init__(self):
        QMainWindow.__init__(self)
        loadUi("power_planner.ui", self)
        self.setWindowTitle("POWER PLANNER 1.0")
        #self.setFixedSize(QSize(1450, 890))
        self.setWindowFlags(QtCore.Qt.Window |
                           QtCore.Qt.CustomizeWindowHint |
                           QtCore.Qt.WindowMinimizeButtonHint |
                           QtCore.Qt.WindowMaximizeButtonHint |
                           QtCore.Qt.WindowCloseButtonHint )

        # DPI에 따라 UI 요소 크기 조정
        self.adjust_ui_scaling()

        # 전역변수 정의
        self.customer_name = ""    # 고객명
        self.customer_no = ""      # 고개번호
        self.contract_kind = ""    # 펌프 데이터파일의 파일경로 및 이름을 저장
        self.contract_capa = None  # 계약용량
        self.peak_power = None     # 피크전력
        self.average_power = None  # 평균전력
        self.min_power = None      # 최소기본전력
        self.gen_power = None      # 발전가능량

        # 기본 디렉토리 
        self.root_dir = None       # 기본 디렉토리 (프로그램이 실행된 디렉토리)


        # 그래프의 x.y 좌표
        self.max_x = 366           # x축 최대 좌표 (초기화)
        self.max_y = 5000          # y축 최대 좌료 (초기화)

        # 2. 기본 디렉토리 Setup 및 만들기
        self.root_dir = os.getcwd()        # 기본 디렉토리를 가져와 설정
        self.directory_setup("planner")    # planner 디렉토리가 없으면 만든다.
        self.directory_setup("pdf")        # pdf 디렉토리가 없으면 만든다.
        self.directory_setup("data")       # data 디렉토리가 없으면 만든다.

        # 테이블 및 그래프 초기화면 설정
        self.setup_table()         # 테이블 설정
        self.plot_graph_start()    # 그래프 화면 초기화화


        # 버튼 정의하는 함수를 실행
        self.setup_buttons()

        # 메뉴 설정
        self.setup_menu()

        # 초기값 데이터 블러오기
        self.open_setup_data()

    def adjust_ui_scaling(self):
        # UI 요소의 크기를 DPI에 맞게 조정
        dpi = QApplication.primaryScreen().logicalDotsPerInch()
        scale_factor = dpi / 96  # 96은 기본 DPI

        # LineEdit 크기 조정
        for line_edit in self.findChildren(QLineEdit):
            line_edit.setFixedHeight(int(25 * scale_factor))  # 높이 조정

        # Label 크기 조정
        for label in self.findChildren(QLabel):
            if label.objectName() != "label_2":  # Label_2를 제외
                font_size = int(10 * scale_factor)
                label.setFont(QFont("Arial", font_size))
        # QPushButton 크기 조정
        for button in self.findChildren(QPushButton):
            button.setFixedHeight(int(30 * scale_factor))  # 높이 조정
            font_size = int(10 * scale_factor)
            button.setFont(QFont("Arial", font_size))  # 폰트 크기 조정

        # 그래프 크기 조정
        #self.MplWidget.setFixedSize(int(1100 * scale_factor), int(600 * scale_factor))  # 원하는 크기로 조정

        # 기타 요소의 크기 조정이 필요할 경우 추가
    def resizeEvent(self, event):
        # Qt Designer에서 정의한 Layer의 크기에 맞추어 그래프 크기 조정
        layer = self.findChild(QWidget, "MplWidget")  # "layerName"을 실제 레이어의 객체 이름으로 변경
        if layer:
            self.MplWidget.setGeometry(layer.geometry())  # 레이어의 크기와 위치에 맞추어 설정
        super().resizeEvent(event)


    def directory_setup(self,directory):
        # 디렉토리가 존재하는지 확인
        if not os.path.exists(directory):
            # 디렉토리가 없으면 생성
            os.mkdir(directory)


    def open_setup_data(self):
        # 모타 데이터 불러오기기
        file_name = "motor_eff.csv"
        #self.pump_motor_data = pd.read_csv(file_name)

    def setup_menu(self):
        # 상단 메뉴를 만든다.
        # File 메뉴 : Open / Exit
        # Power Planner : Power Planner Data Downnload / Calculation 
        # Help : https://pp.kepco.co.kr/

        # 메뉴 바 생성
        menubar = self.menuBar()

        # [File 메뉴]
        file_menu = menubar.addMenu("File")

        # New Data 하위 메뉴 추가
        # Power Planner 데이터 열기
        power_open_action = QAction("Open", self)
        power_open_action.triggered.connect(self.open_power_data)  # Pump Data 액션에 연결
        file_menu.addAction(power_open_action)

        # 프로그램 종료 액션 추가
        exit_action = QAction("End", self)
        exit_action.triggered.connect(self.close_application)  # 종료 액션에 연결
        file_menu.addAction(exit_action)

        # [Calculation 메뉴]
        calculation_menu = menubar.addMenu("Calculation")
        # Power Pla 하위 메뉴 추가
        self.calculation_action = QAction("Calculation", self)
        self.calculation_action.triggered.connect(self.power_calculation)  # Pump Data 액션에 연결
        calculation_menu.addAction(self.calculation_action)
        # [레포트 만들기]
        self.report_action = QAction("Report", self)
        self.report_action.triggered.connect(self.power_report)  # Pump Data 액션에 연결
        calculation_menu.addAction(self.report_action)

        # 초기상태 비활성화
        self.calculation_action.setEnabled(False)
        self.report_action.setEnabled(False)

        # [Power Planner 메뉴]
        power_menu = menubar.addMenu("Power Planner")
        # Power Planner 하위 메뉴 추가
        file_data_action = QAction("Power Planner Data Download", self)
        file_data_action.triggered.connect(self.download_data_window)  # 다운로드 데이터 윈도우를 연다
        power_menu.addAction(file_data_action)

        # Help메뉴 생성
        help_menu = menubar.addMenu("Help")
        # 펌프데이터 열기
        help_action = QAction("https://pp.kepco.co.kr/", self)
        help_action.triggered.connect(self.open_help_link)  # 링크열기 액션에 연결
        help_menu.addAction(help_action)

    def open_help_link(self):
        # 웹사이트 열기
        QDesktopServices.openUrl(QUrl("https://pp.kepco.co.kr/"))

    def close_application(self):
        # 종료 확인 메시지 박스
        reply = QMessageBox.question(self, '종료', '프로그램을 종료하시겠습니까?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.close()  # 프로그램 종료

    def setup_buttons(self):
        #**********************************
        #  버튼 함수를 정의
        #**********************************
        # 버튼을 생성하려면 여기에 버튼의 이름과 연결함수를 정의
        # 1. 버튼을 누르면 연결된 함수가 실행행

        # 1. 버튼의 정의
        self.pushButton_calculation.clicked.connect(self.power_calculation)  # "계산하기" 버튼을 누르면 power_calculation함수를 실행
        self.pushButton_report.clicked.connect(self.power_report)            # "계산하기" 버튼을 누르면 pump_report함수를 실행

    def power_report(self):
        # 펌프의 모델의 디렉토리를 가져온다.
        model_dir = self.label_model_dir.text()

        # *************************
        # MS WORD에 보고서 작성하기
        # *************************

        # 1.모듈 불러오기
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt
        from docx.oxml.ns import qn

        # 2. MS WORD를 doc로 초기화한다.
        doc = Document()
        
        # 3. MS WORD 작성
        # 3.1 상하좌우 여백 설정
        sections = doc.sections
        for section in sections:
            # section.top_margin = Cm(0.5)
            # section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(2.25)
            section.right_margin = Cm(2.25)

        # 3.2 초기 변수 세팅
        customer_name_rpt = self.customer_name
        customer_no_rpt = self.customer_no
        contract_kind_rpt = self.contract_kind
        contract_capa_rpt = self.contract_capa
        search_year_rpt = self.search_year
        gen_power_rpt = self.gen_power
        peak_power_rpt = self.peak_power
        average_power_rpt = self.average_power
        min_power_rpt = self.min_power


        # 3.3 본문 작성
        # Title 제목 작성
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_s = title.add_run('POWER PLANNER REPORT')
        title_s.bold = True
        title_s.font.name = 'NanumGothicCoding'
        title_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        title_s.font.size = Pt(17)

        # 1) 고객현황
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('1. 고객현황')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)      

        # 1-1) 고객명 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  1) 고객명 : ' + customer_name_rpt)
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 1-2) 고객ID 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  2) 고객ID : ' + customer_no_rpt)
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 1-3) 계약종별 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  3) 계약종별 : ' + contract_kind_rpt)
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 1-4) 계약용량 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  4) 계약용량 : ' + contract_capa_rpt + ' kW')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 
        # 한줄 삽입
        article = doc.add_paragraph()
        article_s = article.add_run(' ')


        # 2) 전력현황
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('2. 전력현황 분석 ' +'(' + str(search_year_rpt) + '년도)')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)   
        # 저장한 그래프를 불러온다.
        now_dir = self.root_dir
        file_path = now_dir+"/data/power_chart.png"
        doc.add_picture(file_path, width=Cm(17), height=Cm(8))

        # 2.1) 전력사용 현황 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  2.1 전력 사용현황')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 2.1-1) 피크전력 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('    1) 피크전력 : ' + str(peak_power_rpt) + ' kW')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 2.1-2) 평균수요 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('    2) 평균소요 : ' + str(average_power_rpt) + ' kW')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 2.1-3) 최소기본 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('    3) 최소기본 : ' + str(min_power_rpt) + ' kW')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 2.2) 발전가능량
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('  2.2 발전가능량')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 2.1-1) 발전가능량 분석
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('    1) 발전가능량 : ' + str(gen_power_rpt) + ' kW')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)

        doc.add_page_break()  # 여기서 페이지가 나뉩니다

        # 3) 단가검토
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('3. 단가검토' )
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10) 

        # 3.1) 계절별 발전단가 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('3.1 계절별 발전단가 ')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)

        # 운전점 데이터를 읽어 표로로 나타낸다.
        file_path = "./data/" +  customer_no_rpt + "_" + str(search_year_rpt) + "_" + "unit_price.csv"
        data = pd.read_csv(file_path, encoding='euc-kr', dayfirst=True, parse_dates=[0])
        # 운전점 자료를 리스트 형태로 data에 저장한다.
        data = data.values.tolist()

        # 표 추가
        '''
        표는 4열로 만들고 
        행은 처음 1개행은 Heading 강제로 만들고
        추가 행은 List에서 순차적으로 data를 읽어 추가한다. 
        '''
        table = doc.add_table(rows=1, cols=4)
        table.style = doc.styles['Table Grid']

        # 테이블의 헤딩을 입력한다.
        row = table.rows[0].cells
        row[0].text = '구 분'
        row[1].text = 'SUMMER'
        row[2].text = 'INTERMID'
        row[3].text = 'WINTER'

        # 가운데 정렬
        for cell in row:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 테이블의 행을 순차적으로 List에서 읽어와 추가한다.
        for data_1, data_2, data_3, data_4  in data:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            if data_1 != 0.0 :
                row[0].text = str(data_1)  
                row[1].text = str(data_2)
                row[2].text = str(data_3)
                row[3].text = str(data_4)
                #row[4].text = str(data_5)
            else :
                row[0].text = ""  
                row[1].text = "" 
                row[2].text = "" 
                row[3].text = "" 
                #row[4].text = "" 
                #row[5].text = ""

            # 가운데 정렬
            for cell in row:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 각 열의 너비 설정
        # 각 열에 대해 셀의 너비를 설정합니다.
        table.columns[0].width = Cm(3)  # 운전점
        table.columns[1].width = Cm(2)  # 유량
        table.columns[2].width = Cm(2)  # 양정
        table.columns[3].width = Cm(2)  # 효율
        #table.columns[4].width = Cm(2)  # 동력
        #table.columns[5].width = Cm(2)  # 동력
        #table.columns[5].width = Cm(2)  # 비고

        # 테이블의 글자 크기 변경
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(10)

        # 한줄 삽입
        article = doc.add_paragraph()
        article_s = article.add_run(' ')

        # 3.1) 계절별 발전단가 
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY 
        article_s = article.add_run('3.2 월별 사용요금 분석')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)

        # 운전점 데이터를 읽어 표로로 나타낸다.
        file_path = "./data/" +  customer_no_rpt + "_" + str(search_year_rpt) + "_" + "monthly_price.csv"
        data = pd.read_csv(file_path, encoding='euc-kr', dayfirst=True, parse_dates=[0])
        data['계약전력(kW)'] = data['계약전력(kW)'].apply(lambda x: f'{int(x):,}' if pd.notna(x) and isinstance(x, (int, float)) else x) # 1000단위마다 쉽표 넣어줌
        data['요금적용전력(kW)'] = data['요금적용전력(kW)'].apply(lambda x: f'{int(x):,}' if pd.notna(x) and isinstance(x, (int, float)) else x) # 1000단위마다 쉽표 넣어줌
        data['사용전력량(kWh)'] = data['사용전력량(kWh)'].apply(lambda x: f'{int(x):,}' if pd.notna(x) and isinstance(x, (int, float)) else x) # 1000단위마다 쉽표 넣어줌
        data['전기요금(원)'] = data['전기요금(원)'].apply(lambda x: f'{int(x):,}' if pd.notna(x) and isinstance(x, (int, float)) else x) # 1000단위마다 쉽표 넣어줌
        data = data.fillna('-')         # nan 을 - 으로 변경
        # 운전점 자료를 리스트 형태로 data에 저장한다.
        data = data.values.tolist()

        # 표 추가
        '''
        표는 6열로 만들고 
        행은 처음 1개행은 Heading 강제로 만들고
        추가 행은 List에서 순차적으로 data를 읽어 추가한다. 
        '''
        table = doc.add_table(rows=1, cols=6)
        table.style = doc.styles['Table Grid']

        # 테이블의 헤딩(가로 제목)을 입력한다.
        row = table.rows[0].cells
        row[0].text = '년/월'
        row[1].text = '계약전력(kW)'
        row[2].text = '요금적용전력(kW)'
        row[3].text = '사용전력량(kWh)'
        row[4].text = '전기요금(원)'
        row[5].text = '평균단가(원/kwh)'

        # 가운데 정렬
        for cell in row:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        #print(data)

        # 테이블의 행을 순차적으로 List에서 읽어와 추가한다.
        for _, data_1, data_2, data_3, data_4, data_5, data_6  in data:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            if data_1 != 0.0 :
                row[0].text = str(data_1)  
                row[1].text = str(data_2)
                row[2].text = str(data_3)
                row[3].text = str(data_4)
                row[4].text = str(data_5)
                row[5].text = str(data_6)
            else :
                row[0].text = ""  
                row[1].text = "" 
                row[2].text = "" 
                row[3].text = "" 
                row[4].text = "" 
                row[5].text = ""

            # 가운데 정렬
            for cell in row:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 각 열의 너비 설정
        # 각 열에 대해 셀의 너비를 설정합니다.
        table.columns[0].width = Cm(1)  
        table.columns[1].width = Cm(2)  
        table.columns[2].width = Cm(2)  
        table.columns[3].width = Cm(2)  
        table.columns[4].width = Cm(2) 
        table.columns[5].width = Cm(2) 

        # 테이블의 글자 크기 변경
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(10)

        # 한줄 삽입
        article = doc.add_paragraph()
        article_s = article.add_run(' ')

        # 5) Note
        article = doc.add_paragraph()
        article.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        article.paragraph_format.line_spacing = 1.1
        article.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        article_s = article.add_run('4. Note ')
        article_s.bold = False
        article_s.font.name = 'NanumGothicCoding'
        article_s._element.rPr.rFonts.set(qn('w:eastAsia'), 'NanumGothicCoding')
        article_s.font.size = Pt(10)

        # 3.4 저장하기
        # power_planner.docx로 파일을 저장한다.
        now_dir = self.root_dir
        doc_dir = os.path.join(now_dir, "data", "power_planner.docx")
        #doc_dir = now_dir+"/data/pump_curve.docx"
        doc.save(doc_dir)

        # 9) PDF 파일로 변환하기기
        # doc 파일을 pdf 파일로 변환한다.       
        file_name = "./planner/"+ str(customer_name_rpt) + "_" + str(customer_no_rpt) + "_" + str(search_year_rpt) + ".pdf"
        file_name_1, _ = os.path.splitext(os.path.basename(file_name))
        try:
            pdf_file_path = os.path.join(now_dir, "pdf", f"{file_name_1}.pdf")
            convert(doc_dir, pdf_file_path)  # doc_dir에서 pdf_file_path로 변환
            # 레포트 생성 메시지
            QMessageBox.information(self, 'Report', 'Report가 생성되었습니다.')
        except Exception as e:
            QMessageBox.critical(self, 'Error', f'Error가 발생하였습니다.: {str(e)}{pdf_file_path}')

    def setup_table(self):
        #**********************************
        #  CRT에 있는 테이블의 크기를 설정
        #**********************************
        # 1. Table Widget에 있는 열의 크기를 정의
        # 2. 테이블의 첫번재 줄의 제목을 가운데 정열

        # 1.각 열의 크기를 설정
        self.price_table.setColumnWidth(0, 150)  # 첫 번째 열 너비를 150px로 설정
        self.price_table.setColumnWidth(1, 150)  # 첫 번째 열 너비를 150px로 설정
        self.price_table.setColumnWidth(2, 150)  # 첫 번째 열 너비를 150px로 설정
        self.price_table.setColumnWidth(3, 150)  # 첫 번째 열 너비를 100px로 설정
        #self.price_table.setColumnWidth(4, 150)  # 첫 번째 열 너비를 150px로 설정
        # 2. 가운데 정열
        
        self.price_table.item(0, 0).setTextAlignment(Qt.AlignVCenter | Qt.AlignRight)



    def open_power_data(self):
        # [함수설명]
        # CSV데이터 형대로 저장된 Power Planner의 데이터를 불러와서 CRT에 뿌려줍니다.
        # 메뉴 활성화
        self.calculation_action.setEnabled(True)
        self.report_action.setEnabled(False)
        # 버튼 활성화
        self.pushButton_calculation.setEnabled(True)
        self.pushButton_report.setEnabled(False)
        #**********************************
        #  저장된 Power Planner 데이터 불러오기
        #**********************************
        # 기존 저장된 펌프 데이터를 불러옵니다.
        
        # 1. 파일 Dialog를 이용하여 파일을 불러온다.
        # 1.1 QFileDialog를 사용하여 파일 선택
        # 파일을 root_dir로 지정하여 QFileDialog를 읽는다.
        file_name, _ = QFileDialog.getOpenFileName(self, "Select Power Data", self.root_dir, "data Files (*.pwr)")
        if not file_name:
            return  # 파일이 선택되지 않은 경우 함수 종료
        
        # 1.2 선택한 파일을 읽어 DataFrame으로 변환
        user_data_df = pd.read_csv(file_name,encoding='euc-kr',dtype=str)
        #print(user_data_df)
        #print(file_name)

        # 2. 초기 그래프를 그린다.
        self.plot_graph_start()

        # 3. DataFrame의 user data를 추출
        # 3.1 각 펌프데이터 추출하여 변수에 입력
        customer_name_1 = user_data_df.iloc[0:1,1:2].values[0][0]     # 고객명
        self.customer_name = customer_name_1
        #print(user_data_df)     
        customer_no_1 = str(user_data_df.iloc[0:1,2:3].values[0][0])       # 고객번호
        self.customer_no = customer_no_1
        #print(customer_no_1)
        customer_pw_1 = user_data_df.iloc[0:1,3:4].values[0][0]       # 고객비밀번호
        #print(customer_pw_1)
        contract_kind_1 = user_data_df.iloc[0:1,4:5].values[0][0]     # 계약종류
        self.contract_kind = contract_kind_1
        #print(contract_kind_1)
        contract_capa_1 = user_data_df.iloc[0:1,5:6].values[0][0]     # 계약용량
        self.contract_capa = contract_capa_1
        #print(contract_capa_1)
        search_year_1 = user_data_df.iloc[0:1,6:7].values[0][0]       # 검색년도
        self.search_year = search_year_1
        #self.search_year = search_year_1
        #print(search_year_1)
        
        #**********************************
        #  읽은 데이터를 CRT에 표기하기
        #**********************************
        # 기존 저장된 펌프 데이터를 CRT에 나타냅니다.

        # 1. 고객정보를 CRT에 표기하기
        # 1.1 고객명    
        customer_name = self.findChild(QLineEdit, "customer_name_scr")
        if customer_name:
            if customer_name_1 != "" and not (isinstance(customer_name_1, float) and math.isnan(customer_name_1)):
                customer_name.setText(str(customer_name_1))
            else :
                customer_name.setText("")  # 빈 문자열로 설정
            customer_name.setAlignment(Qt.AlignCenter)  # 오른쪽 정렬
        # 1.2 고객번호
        customer_no = self.findChild(QLineEdit, "customer_no_scr")
        if customer_no:
            if customer_no_1 != "" and not (isinstance(customer_no_1, float) and math.isnan(customer_no_1)):
                customer_no.setText(str(customer_no_1))
            else :
                customer_no.setText("")  # 빈 문자열로 설정
            customer_no.setAlignment(Qt.AlignCenter)  # 오른쪽 정렬
        # 1.3 비밀번호
        search_year = self.findChild(QLineEdit, "search_year_scr")
        if search_year:
            if search_year_1 != "" and not (isinstance(search_year_1, float) and math.isnan(search_year_1)):
                search_year.setText(str(search_year_1))
            else :
                search_year.setText("")  # 빈 문자열로 설정
            search_year.setAlignment(Qt.AlignCenter)  # 오른쪽 정렬

    def plot_graph_start(self):
        # 그래프 초기화 하기

        # 이전 그래프 지우기
        self.MplWidget.canvas.axes.clear()
        
        # 그래프의 x축과 y축 설정
        self.MplWidget.canvas.axes.set_xlim(0, 365 - 1)
        self.MplWidget.canvas.axes.set_ylim(0, 5000 + 1)  # y축 범위 설정

        # 레전드 추가
        self.MplWidget.canvas.axes.set_xlabel('DAY')
        self.MplWidget.canvas.axes.set_ylabel('Daily Max (KW)')
        self.MplWidget.canvas.axes.grid()                         # 그리드를 표기

        # 그래프 업데이트
        self.MplWidget.canvas.draw()

        self.MplWidget.canvas.figure.tight_layout()  # 여백 조정

    def power_calculation(self):
        # 펌프의 계산하기를 클릭하면  plot_graph() 함수를 
        # 레포트메뉴 활성화
        self.report_action.setEnabled(True)
        # 레포트 버튼 활성화
        self.pushButton_report.setEnabled(True)

        self.plot_graph()

    def plot_graph(self):
        # 계산하기 버튼을 누르면 이 함수가 실행된다.
        
        # 1. 변수할당
        # 1.1 계약전류/계약종류/기준디렉토리
        # 전역변수의 값을 지역변수로 할당
        contract_capa = self.contract_capa    # 계약전력
        contract_kind = self.contract_kind    # 계약종류
        root_dir = self.root_dir              # 기준디렉토리
        #print(root_dir)
        #**********************************
        #  Line Edit의 고객현황 변수에 저장
        #**********************************
        
        # 1.2 우측 팔레트에 고객데이터 출력
        # 1) 고객명
        try:
            customer_name = str(self.customer_name_scr.text())
        except ValueError:
            customer_name = " "
        # 2) 고객번호
        try:
            customer_no = self.customer_no_scr.text()
        except ValueError:
            customer_no = " "
        # 3) 분석년도
        try:
            search_year = self.search_year_scr.text()
        except ValueError:
            search_year = " "

        # 2. 고객정보 외의 자료를 파일에서 읽어온다
        # 2.1 계절별 단가 파일
        file_name = "./data/"+customer_no+"_"+search_year+"_unit_price.csv"
        #print(file_name)
        unit_price= pd.read_csv(file_name,encoding='euc-kr')
        # 2.2 최대전력 사용량 파일
        file_name = "./data/"+customer_no+"_"+search_year+"_max_power_daily.csv"
        max_power_daily_df= pd.read_csv(file_name,encoding='euc-kr')
        # 2.3 일별전력 사용량 파일 
        file_name = "./data/"+customer_no+"_"+search_year+"_daily_power_use.csv"
        daily_power_use_df= pd.read_csv(file_name,encoding='euc-kr')
        
        # 3. 월별 날수 정의
        month_day = [31,28,31,30,31,30,31,31,30,31,30,31]
        
        # 4. 최대전력 Graph를 그리기 위해 DataFrame을 일렬로 정렬
        # 4.1 1월 최대전력 대에터 추출
        data_pday_graph = max_power_daily_df.iloc[0:30, 1]
        # 4.2 2월~12월 최대전력 데이터를 일렬로 정렬
        for p_m in range(2,13) :
            p_d = month_day[p_m-1]
            data_pday_graph = pd.concat([data_pday_graph, max_power_daily_df.iloc[0:p_d,p_m]], axis=0,ignore_index = True)

        # 5. 발전가능량 계산
        # 5.1 계약전력 추출
        max_power = float(contract_capa[:-2])
        # 5.2 발저가능량 계산 (계약전력의 30%)
        exp_gen = data_pday_graph.max() - max_power*0.3

        #**********************************
        # 메인화면 데이터 입력
        #**********************************
        
        # 1. 메인화면 하단 table에 계절별 단가 입력
        # 1.1 기본요금
        self.price_table.setItem(0,0, QTableWidgetItem(str(round(unit_price.iloc[0,1],2)))) 
        self.price_table.item(0, 0).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(0,1, QTableWidgetItem(str(round(unit_price.iloc[0,2],2))))
        self.price_table.item(0, 1).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(0,2, QTableWidgetItem(str(round(unit_price.iloc[0,3],2))))
        self.price_table.item(0, 2).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 

        # 1.2 경부하
        self.price_table.setItem(1,0, QTableWidgetItem(str(round(unit_price.iloc[1,1],2)))) 
        self.price_table.item(1, 0).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(1,1, QTableWidgetItem(str(round(unit_price.iloc[1,2],2))))
        self.price_table.item(1, 1).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(1,2, QTableWidgetItem(str(round(unit_price.iloc[1,3],2))))
        self.price_table.item(1, 2).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 

        # 1.3 중부하
        self.price_table.setItem(2,0, QTableWidgetItem(str(round(unit_price.iloc[2,1],2)))) 
        self.price_table.item(2, 0).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(2,1, QTableWidgetItem(str(round(unit_price.iloc[2,2],2))))
        self.price_table.item(2, 1).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(2,2, QTableWidgetItem(str(round(unit_price.iloc[2,3],2))))
        self.price_table.item(2, 2).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 

        # 1.4 최대부하
        self.price_table.setItem(3,0, QTableWidgetItem(str(round(unit_price.iloc[3,1],2)))) 
        self.price_table.item(3, 0).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(3,1, QTableWidgetItem(str(round(unit_price.iloc[3,2],2))))
        self.price_table.item(3, 1).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 
        self.price_table.setItem(3,2, QTableWidgetItem(str(round(unit_price.iloc[3,3],2))))
        self.price_table.item(3, 2).setTextAlignment(Qt.AlignHCenter | Qt.AlignVCenter)  # 가운데 정렬 

        # 2. 메인화면 우측에 년간 전력사용현황 입력
        # 2.1 평균값을 계산하기 위하여 기술통계값을 계산
        desc=data_pday_graph.describe()
        #print(desc)
        # 2.2 오른쪽 년간 전력사용현황에 값을 입력
        self.contract_kind_scr.setText(str(contract_kind))
        self.contract_capa_scr.setText(str(contract_capa[0:-2]))
        self.peak_power_scr.setText(str(round(data_pday_graph.max(),1)))
        self.average_power_scr.setText(str(round(desc[1],1)))
        self.base_min_scr.setText(str(round(max_power*0.3,1)))

        # 계산 후 레포트 작성용 변수 저장
        self.peak_power = round(data_pday_graph.max(),1)
        self.average_power = round(desc[1],1)
        self.min_power = round(max_power*0.3,1)
        self.gen_power = round(exp_gen,1)


        #**********************************
        #        그래프를 그린다.
        #**********************************

        # 1. 그래프 한글폰트 정의
        path ="c:/Windows/Fonts/malgun.ttf"
        font_name = font_manager.FontProperties(fname=path).get_name()
        rc('font', family = font_name)

        # 2. 그래프를 초기화 한다.
        # 2.1 그래프 초기화
        self.MplWidget.canvas.axes.clear()
        graph_start = 0
        max_x = self.max_x
        max_y = max_power *1.05
        # 2.2 그래프의 x축과 y축 설정
        self.MplWidget.canvas.axes.set_xlim(graph_start, max_x - 1)     # x축 범위 설정
        self.MplWidget.canvas.axes.set_ylim(0, max_y)                    # y축 범위 설정

        # 3. 그래프 그리기
        # 3.1 일별 최대수요 그래프 데이터 정리
        data_pday_graph = data_pday_graph.reset_index(drop=True)
        # 3.2 주석 값 설정
        exp_gen_text = "발전가능용량 : " + str(round(exp_gen,0))+"kW"
        year_text = "기준년도 : " + str(search_year)+"년"
        c_name_text = "고객명 : " + customer_name
        power_level = "계약종별 : " + contract_kind
        # 3.3 그래프 그리기
        # 1) 최대수요 그래프 그리기
        self.MplWidget.canvas.axes.plot(data_pday_graph.index, data_pday_graph,color='red', linestyle="-", label='최대수요(kW)', marker='')
        # 2) 분석자료(계약전력, 피크전력, 평균수요, 최소기본전력) 직선 그리기
        self.MplWidget.canvas.axes.axhline(y=max_power, color='magenta', linestyle='--', label ="계약전력 $(%.1f$kW)" %(max_power))
        self.MplWidget.canvas.axes.axhline(y=exp_gen+max_power*0.3, color = 'red', linestyle =':', label ="Peak 전력 $(%.1f$kW)" %(exp_gen+max_power*0.3))
        self.MplWidget.canvas.axes.axhline(y=desc[1], color='green', linestyle='--', label ="평균수요 $(%.1f$kW)" %(desc[1]))
        self.MplWidget.canvas.axes.axhline(y=max_power*0.3, color='blue', linestyle='--', label ="최소기본전력 $(%.1f$kW)" %(max_power*0.3))
        # 3) 주석표기
        self.MplWidget.canvas.axes.text(10,max_power*0.95,c_name_text)
        self.MplWidget.canvas.axes.text(10,max_power*0.92,year_text)
        self.MplWidget.canvas.axes.text(10,max_power*0.89,exp_gen_text)
        self.MplWidget.canvas.axes.text(10,max_power*0.86,power_level)
        # 4) 레전드 추가
        self.MplWidget.canvas.axes.legend(loc='best')             # 최적위치에 Legend를 그린다.
        self.MplWidget.canvas.axes.set_xlabel('DAY')              # 가로축 Legend
        self.MplWidget.canvas.axes.set_ylabel('최대수요(KW)')      # 새로축 Legend
        self.MplWidget.canvas.axes.grid()                         # 그리드를 표기
        # 5) 그래프 업데이트
        self.MplWidget.canvas.draw()

        # 4. 그래프 저장
        now_dir = self.root_dir
        file_path = now_dir+"/data/power_chart.png"
        #file_path = "pump_chart.png"  # 저장할 파일 경로 및 이름
        self.MplWidget.canvas.figure.tight_layout()  # 여백 조정
        self.MplWidget.canvas.figure.savefig(file_path)  # 그래프 저장

        

    def download_data_window(self):
        # 모델을 만들기 위한 새로운 윈도우를  생성한다.
        self.power_data_window = QDialog(self)
        loadUi("power_data.ui", self.power_data_window)  # pump_data.ui 파일 로드, power_data_window를 로드시킨ㄷ.
        self.power_data_window.setWindowTitle("Data Download System")
        self.power_data_window.setFixedSize(730, 285)

        # QLineEdit 초기화
        self.customer_name = self.power_data_window.findChild(QLineEdit, "customer_name_scr")
        self.customer_no = self.power_data_window.findChild(QLineEdit, "customer_no_scr")
        self.customer_pw = self.power_data_window.findChild(QLineEdit, "customer_pw_scr")
        self.search_year = self.power_data_window.findChild(QLineEdit, "search_year_scr")

        # Download 버튼 클릭 시 power_planner_login 함수 수행
        self.login_btn_1 = self.power_data_window.findChild(QPushButton, "download_btn")  # 버튼을 누르면 copy_table_1 함수를 실행
        self.login_btn_1.clicked.connect(self.power_planner_login)

        # 확인 버튼 클릭 시 현재 윈도우 닫음
        self.ok_button = self.power_data_window.findChild(QPushButton, "ok_btn")  # "ok_btn"을 누르면 메인화면으로 이동
        self.ok_button.clicked.connect(self.power_data_window.accept) 

        # 새로운 윈도우 표시
        self.power_data_window.exec_()


    def power_planner_login(self):

        #==============================
        # 한전 파워플래너에 Login
        #==============================
        chrome_options = Options()
        chrome_options.add_experimental_option("excludeSwitches", ["enable-logging"])

        # 1. power_data.ui의 customer_no_scr과 customer_pw_scr, search_year_scr에서 고객번호와 비밀번호, 분석년도를 가져온다.
        user_name = self.customer_name.text()          # 고객번호
        user_id = self.customer_no.text()              # 고객번호
        user_pw = self.customer_pw.text()              # 고객비밀번호
        search_year = int(self.search_year.text())     # 분석년도

        # 2. login 페이지를 열어 ID와 PW를 입력한다.
        # 2.1 login 하고자 하는 login page를 url에 저장한다. (1)
        url = "https://pp.kepco.co.kr/intro.do"
        browser = webdriver.Chrome(options=chrome_options)
        # 2.2 Overlay되어 가려질 것을 대비하여 선재적 기본브라우져의 크기를 크게한다.
        browser.set_window_size(1400, 1000) # 가림 방지용 여유 뷰포트

        # 2.3 공지/배너/오버레이 제거한다. (notice_auto_cont 기준)
        try:
            overlay = browser.find_element(By.ID, "notice_auto_cont") 
            try:
                close_btn = overlay.find_element(By.CSS_SELECTOR, ".close, .btn_close, [aria-label='닫기']")
                close_btn.click()
                wait.until(EC.invisibility_of_element_located((By.ID, "notice_auto_cont")))
            except Exception:
                browser.execute_script("arguments[0].style.display='none';", overlay)
                time.sleep(0.1)
        except Exception:
            pass

        # 2.4 3초 대기하기 ----(3)
        browser.implicitly_wait(5)
        # 2.5 로그인하기   ---- (4)
        browser.get(url)
        element_id = browser.find_element('id','RSA_USER_ID')
        element_id.clear()
        element_id.send_keys(user_id)
        element_pw = browser.find_element('id','RSA_USER_PWD')
        element_pw.clear()
        element_pw.send_keys(user_pw)

        # 3. 버튼누르기(5)
        browser.find_element('xpath','//*[@id="intro_form"]/form/fieldset/input[1]').click()

        time.sleep(5)
        
        # 로그인 에러시 메시지 팝업
        # 로그인 성공 시 URL은 intro.do가 아닐 것이라고 가정합니다.
        if "intro.do" in browser.current_url.lower(): # 이 'if' 문 블록입니다.
                QMessageBox.warning(self, "로그인 실패", "고객번호 또는 비밀번호가 올바르지 않습니다. 다시 확인해 주세요. ㅠㅠ")
                # 브라우저 종료 로직
                if browser is not None:
                    try:
                        browser.quit()
                    except Exception as quit_e:
                        print(f"브라우저 종료 중 오류 발생: {quit_e}")
                return 

        #==============================
        # 고객정보를 가져온다.
        #==============================

        # 1. 고객정보를 분석한다.
        # 1.1 고객정보 웹페이지를 url에 저장한다. (1)
        url="https://pp.kepco.co.kr/mb/mb0101.do?menu_id=O010601"
        browser.set_window_size(1400, 1000) # 가림 방지용 여유 뷰포트

        # 1.2 url 주소로 가상브라우져를 뛰운다  (2)
        browser.get(url)

        # 공지/배너/오버레이 제거 시도(notice_auto_cont 기준)
        try:
            overlay = browser.find_element(By.ID, "notice_auto_cont") 
            try:
                close_btn = overlay.find_element(By.CSS_SELECTOR, ".close, .btn_close, [aria-label='닫기']")
                close_btn.click()
                wait.until(EC.invisibility_of_element_located((By.ID, "notice_auto_cont")))
            except Exception:
                browser.execute_script("arguments[0].style.display='none';", overlay)
                time.sleep(0.1)
        except Exception:
            pass

        # 1.3 class name이 "table_info"인 부분을 찾아 pre_elem에 저장한다. (3)
        pre_elem = browser.find_element(By.CLASS_NAME,"table_info")

        # 1.4 pre_elem내에서 "tbody"라는 tag name을 찾아 client_name에 저장한다. (4)
        client_name = pre_elem.find_element(By.TAG_NAME,"tbody")

        # 1.5 client_name의 내용중 공백으로 된 부분을 분리하여 리스트로 만든다. (5)
        client_name = client_name.text.split(" ")
        
        # (데이터확인시) client_name 리스트중 고객번호는 1번, 계약종별은 4번, 계약전력은 6번 이므로 이들을 찾아서 프린트한다 . (6)
        #print("1) 고객이름 : ", user_name)
        #print("1) 고객번호 : ", client_name[1])
        #print("2) 계약종별 : ", client_name[4])
        #print("3) 계약전력 : ", client_name[6])
        # 1.6 데이터를 딕션어리 형태로 Dataframe으로 만든 후 csv로 저장한다. (7)
        user_data_dict = {
            "User Name": [user_name],       # <<-- 값을 리스트로 묶습니다
            "User ID": [user_id],           # <<-- 값을 리스트로 묶습니다
            "User Password": [user_pw],     # <<-- 값을 리스트로 묶습니다
            "Contract Kind" : [client_name[4]],
            "Contract Capa" : [client_name[6]],
            "Search Year": [search_year]    # <<-- 값을 리스트로 묶습니다
        }
        user_data_df = pd.DataFrame(user_data_dict)
        file_name = "./planner/"+ user_data_df["User Name"].iloc[0]+"_"+str(search_year)+"_"+ user_data_df["User ID"].iloc[0] + ".pwr"
        user_data_df.to_csv(file_name,encoding='euc-kr')

            # 데이터 로그린 완료 !!!


        #==============================
        # 계절별/부하대별 전력단가 가져오기
        #==============================

        # 1. 스마트 뷰로 들어가 전력단가를 가져온다.

        # 1.1 스마트뷰 웹페이지를 url에 저장한다.(1)
        url = "https://pp.kepco.co.kr/rm/rm0101.do?menu_id=O010101" 
        browser.set_window_size(1400, 1000) # 가림 방지용 여유 뷰포트

        # 1.2 url 주소로 가상 브라우져를 뛰운다. (2)
        browser.get(url)

        # 6) 공지/배너/오버레이 제거 시도(notice_auto_cont 기준)
        try:
            overlay = browser.find_element(By.ID, "notice_auto_cont") 
            try:
                close_btn = overlay.find_element(By.CSS_SELECTOR, ".close, .btn_close, [aria-label='닫기']")
                close_btn.click()
                wait.until(EC.invisibility_of_element_located((By.ID, "notice_auto_cont")))
            except Exception:
                browser.execute_script("arguments[0].style.display='none';", overlay)
                time.sleep(0.1)
        except Exception:
            pass

        try :
            # 1.3 class_name이 "free_wrap"인 부분을 찾아 pre_elem에 저장한다. (3)
            pre_elem = browser.find_element(By.CLASS_NAME,"fee_wrap")
            #pre_elem = browser.find_element(By.CLASS_NAME,"fee_wrap")
            
            unit_price_raw = pre_elem.text.replace("\n", " ").split(" ")
            #print("unit_price",unit_price_raw)

            # 2. 전력요금 단가를 정리한다.

            # 2.1 unit_price_raw[0]이 전력단가이므로 전력단가내의 쉽표(,) 부분을 삭제한다. (1)
            base_price = float(unit_price_raw[10].replace(",",""))

            # 2.2 Data Frame 기본형을 Dic 형태로 만들어 준다. (1)
            unit_price = {'summer': [0, 0, 0, 0],
                    'intermid' : [0, 0, 0, 0],
                    'winter'   : [0, 0, 0, 0]}

            # 2.3 Dic형태의 data를 DataFrame으로 변경한다. (2)
            unit_price = pd.DataFrame(unit_price)
            unit_price.index = ["기본요금", " 경부하", "중부하", "최대부하"]

            # 2.4 기본형 Data Frame에 계절 및 경부하,중부하,최대부하별로 각각 입력한다. (3)
            unit_price.iloc[0,0] = base_price          #summer, 기본요금
            unit_price.iloc[1,0] = unit_price_raw[12]   #summer, 경부하
            unit_price.iloc[2,0] = unit_price_raw[16]   #summer, 중부하
            unit_price.iloc[3,0] = unit_price_raw[20]  #summer, 최대부하

            unit_price.iloc[0,1] = base_price          #intermid, 기본요금
            unit_price.iloc[1,1] = unit_price_raw[13]   #intermid, 경부하
            unit_price.iloc[2,1] = unit_price_raw[17]   #intermid, 중부하
            unit_price.iloc[3,1] = unit_price_raw[21]  #intermid, 최대부하

            unit_price.iloc[0,2] = base_price          #winter, 기본요금
            unit_price.iloc[1,2] = unit_price_raw[14]   #winter, 경부하
            unit_price.iloc[2,2] = unit_price_raw[18]   #winter, 중부하
            unit_price.iloc[3,2] = unit_price_raw[22]  #winter, 최대부하

            # 3. 전력단가를 저장
            file_name = "./data/"+ user_data_df["User ID"].iloc[0]+"_"+str(search_year)+"_unit_price.csv"
            unit_price.to_csv(file_name,encoding='euc-kr')
        except :
            QMessageBox.information(self,'Error Message','전력단가를 가져오는중에 에러가 발생하였습니다.')
            browser.quit()
            return


        #==============================
        # 최대수요 전력량 다운로드
        #==============================

        # 1. 최대수요량을 다운받는다
        # 1.1 년도와 월의 선택년도/월의 번호를 초기화 한다. (마우스로 선택한 효과를 나타내기 위해 xpath를 이용할 계획) 
        y_i = 0    # 년도의 xpath의 년도순서 번호를 초기화
        m_i = 0    # 월의 xpath의 년도순서번호를 초기화
        y_i = search_year - 2010 +1  # 2010년을 기준으로 년도순서를 계산한다. (한전파워플래너의 메뉴의 첫번째 년도가 2010으로 되어 있음)

        # 1.2 각 월의 일수를 초기화 한다. (data_pdata_graph의 data를 일렬로 세워 그래프를 그릴 경우 사용한다.)
        month_day = [31,28,31,30,31,30,31,31,30,31,30,31]

        # 2. 최대수요를 스크래핑 한다.
        # 2.1 실시간사용량>일별 웹페이지를 url에 저장한다.(1)
        url = "https://pp.kepco.co.kr/rs/rs0102.do?menu_id=O010202" 
        browser.set_window_size(1400, 1000) # 가림 방지용 여유 뷰포트

        # 2.2 url 주소로 가상 웹브라우져를 뛰운다  (2)
        browser.get(url)
        
        # 6) 공지/배너/오버레이 제거 시도(notice_auto_cont 기준)
        try:
            overlay = browser.find_element(By.ID, "notice_auto_cont") 
            try:
                close_btn = overlay.find_element(By.CSS_SELECTOR, ".close, .btn_close, [aria-label='닫기']")
                close_btn.click()
                wait.until(EC.invisibility_of_element_located((By.ID, "notice_auto_cont")))
            except Exception:
                browser.execute_script("arguments[0].style.display='none';", overlay)
                time.sleep(0.1)
        except Exception:
            pass

        # 2.3 url이 loading 될 동안 잠시 기다린다 (3)
        time.sleep(5)


        try: 

            # 2.4 가상 웹브라우져에서 년도, 월, 1일기준, 조회버튼을 xpath를 이용하여 클릭한다.
            
            # [참고사항]
            # 1) 년도버튼     xpath 주소  : //*[@id="SEARCH_YEAR"]
            # 2) 년도선택버튼 xpath 주소  : /html/body/div[2]/div[3]/div[2]/div/p[1]/select[1]/option[1]
            # 3) 월  버튼     xpath 주소  : //*[@id="SEARCH_MONTH"]
            try: # 마우스를 화면 좌상단 빈 곳으로 이동해 hover 해제 
                ActionChains(browser).move_by_offset(-10000, -10000).perform() 
                time.sleep(0.1) 
            except Exception: 
                pass


            for m_i in range(1,13,1):
                browser.find_element(By.XPATH,'//*[@id="SEARCH_YEAR"]').click() # 년도 클릭
                #browser.implicitly_wait(1)
                browser.find_element(By.XPATH,'/html/body/div[2]/div[3]/div[2]/div/p[1]/select[1]/option['+str(y_i)+']').click() # 2017년 클릭

                browser.find_element(By.XPATH,'//*[@id="SEARCH_MONTH"]').click() # 월 클릭
                browser.find_element(By.XPATH,'/html/body/div[2]/div[3]/div[2]/div/p[1]/select[2]/option['+str(m_i)+']').click() # 1월 클릭

                browser.find_element(By.XPATH,'//*[@id="NEEDLE"]').click() # 1일기준 클릭 

                browser.find_element(By.XPATH,'//*[@id="txt"]/div[2]/div/p[2]/span[1]/a/img').click() #  조희 클릭

                # 브라우져가 작동할 시간을 강제적으로 3초정도 준다.
                time.sleep(5)

                browser.find_element(By.XPATH,'//*[@id="kW"]').click() #  최대수요량

                time.sleep(5)

                html = browser.page_source
                data = pd.read_html(html, header=0)
                #print("data_1",data[1])
                #print("data_2",data[2])
                #print("data_3",data[3])
                #print("data_4",data[4])
                # 20250923 홈페이지가 변경되어 리스트가 [3]에서 [4]로 변경됨

                # 추출한 대이터를 정렬한다.
                if (m_i == 1) :
                    data_15_pday_raw = data[4]
                    data_15_pday_1 =data_15_pday_raw.iloc[:,[1]]
                    data_15_pday_1.columns = ['peak'] 
                    data_15_pday_2 = data_15_pday_raw.iloc[:,[5]]
                    data_15_pday_2.columns = ['peak'] 
                    data_pday = pd.concat([data_15_pday_1, data_15_pday_2], axis=0, ignore_index = True)
                else :
                    data_15_pday_raw = data[4]
                    data_15_pday_1 =data_15_pday_raw.iloc[:,[1]]
                    data_15_pday_1.columns = ['peak'] 
                    data_15_pday_2 = data_15_pday_raw.iloc[:,[5]]
                    data_15_pday_2.columns = ['peak'] 
                    data_pday_temp = pd.concat([data_15_pday_1, data_15_pday_2], axis=0, ignore_index = True)
                    data_pday = pd.concat([data_pday, data_pday_temp], axis=1, ignore_index = True)

            # 3. 스크래핑한데이터를 정리(Data Mining)한다.
            #print(data_pday)
            data_pday = data_pday.replace('-','NaN')                       # - 을 NaN으로 변경
            data_pday = data_pday.astype('float')                          # 데이터 타입을 float으로 변경
            data_pday = data_pday.apply(pd.to_numeric,errors='ignore')     # 숫자가 아닌 부분은 무시
            data_pday = data_pday.fillna(0)                                # NaN 부분을 0으로 변경
            data_pday = data_pday.drop(data_pday.index[-1])                # 마지막 행(합계부분)을 삭제
            data_pday.columns = [1,2,3,4,5,6,7,8,9,10,11,12]               # 컬럼명을 1~12월로 변경
            data_pday.index = [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]  # 인덱스를 1~31일로 변경

            # 4. 최대수요량을 저장한다.
            file_name = "./data/"+ user_data_df["User ID"].iloc[0]+"_"+str(search_year)+"_max_power_daily.csv"
            data_pday.to_csv(file_name,encoding='euc-kr')

        except :
            QMessageBox.information(self,'Error Message','최대수요량을 가져오는중에 에러가 발생하였습니다.')
            browser.quit()
            return

        #==============================
        # 일간전력사용량을 다운로드
        #==============================

        # 1. 3초 대기하기
        time.sleep(3)

        #try :
        # 2. 일간전력사용량을 스크래핑 한다.
        for m_i in range(1,13,1):
            # 2.1 년도 클릭
            # 1) 년도버튼을 클릭 
            browser.find_element(By.XPATH,'//*[@id="SEARCH_YEAR"]').click() # 년도 클릭
            # 2) 년도 선택 클릭
            browser.find_element(By.XPATH,'/html/body/div[2]/div[3]/div[2]/div/p[1]/select[1]/option['+str(y_i)+']').click() # 2017년 클릭
            
            # 2.2 월 클릭
            # 1) 월 버튼 클릭
            browser.find_element(By.XPATH,'//*[@id="SEARCH_MONTH"]').click() # 월 클릭
            # 2) 월을 선택 클릭
            browser.find_element(By.XPATH,'/html/body/div[2]/div[3]/div[2]/div/p[1]/select[2]/option['+str(m_i)+']').click() # 1월 클릭
            
            # 2.3 일 클릭
            # 1) 일 선택 클릭
            browser.find_element(By.XPATH,'//*[@id="NEEDLE"]').click() # 1일기준 클릭 
            # 2.3 조회 버튼 클릭
            browser.find_element(By.XPATH,'//*[@id="txt"]/div[2]/div/p[2]/span[1]/a/img').click() #  조희 클릭

            # 2.4 브라우져가 작동할 시간을 강제적으로 3초정도 준다.
            time.sleep(5)
            
            # 2.5 kWh 버튼 클릭 (일간전력사용량)
            browser.find_element(By.XPATH,'//*[@id="kWh"]/a').click() #  일별사용량
            time.sleep(5)
            
            # 2.6 페이지 소스를 가져온다.
            html = browser.page_source
            # 2.7 HTML 테이블을 판다스 데이터프레임으로 변환한다.
            data = pd.read_html(html, header=0)

            #print("data_1",data[1])
            #print("data_2",data[2])
            #print("data_3",data[3])
            #print("data_4",data[4])

            # 2.8 추출한 대이터를 정렬한다.
            # 20250923 홈페이지가 변경되어 리그트[3]이 [4]로 변경됨
            if (m_i == 1) :
                data_15_day_raw = data[4]
                data_15_day_1 =data_15_day_raw.iloc[:,[1]]
                data_15_day_1.columns = ['peak'] 
                data_15_day_2 = data_15_day_raw.iloc[:,[5]]
                data_15_day_2.columns = ['peak'] 
                data_day = pd.concat([data_15_day_1, data_15_day_2], axis=0, ignore_index = True)
            else :
                data_15_day_raw = data[4]
                data_15_day_1 =data_15_day_raw.iloc[:,[1]]
                data_15_day_1.columns = ['peak'] 
                data_15_day_2 = data_15_day_raw.iloc[:,[5]]
                data_15_day_2.columns = ['peak'] 
                data_day_temp = pd.concat([data_15_day_1, data_15_day_2], axis=0, ignore_index = True)
                data_day = pd.concat([data_day, data_day_temp], axis=1, ignore_index = True)

        # 3. Data 처리 (Data Mining)
        data_day = data_day.replace('-','NaN')
        data_day = data_day.astype('float')
        data_day = data_day.apply(pd.to_numeric,errors='ignore')
        data_day = data_day.fillna(0)
        data_day = data_day.drop(data_day.index[-1])
        data_day.columns = [1,2,3,4,5,6,7,8,9,10,11,12]
        data_day.index = [1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26,27,28,29,30,31]

        # 4.일간전력사용량을 저장한다.
        file_name = "./data/"+ user_data_df["User ID"].iloc[0]+"_"+str(search_year)+"_daily_power_use.csv"
        data_day.to_csv(file_name,encoding='euc-kr')

        #except :
        #    QMessageBox.information(self,'Error Message','일간사용전력량을 가져오는중에 에러가 발생하였습니다.')
        #    browser.quit()
        #    return       
        


        #==============================
        # 월별 사용요금 다운로드
        #==============================

        # 1. 월별 사용량 및 사용요금을 다운받는다
        # 1.1 년도와 월의 선택년도/월의 번호를 초기화 한다. (마우스로 선택한 효과를 나타내기 위해 xpath를 이용할 계획) 

        y_i = 2025 + 1 - search_year # 년도 XPath 숫자를 계산  (한전파워플래너의 메뉴의 년도가 2025 - 2019년으로 되어 있음)

        # 2. 월별사용량 및 사용요금을 스크래핑 한다.
        # 2.1 월별청구요금 웹페이지를 url에 저장한다.(1)
        url = "https://pp.kepco.co.kr/cc/cc0102.do?menu_id=O010405" 
        browser.set_window_size(1400, 1000) # 가림 방지용 여유 뷰포트

        # 2.2 url 주소로 가상 웹브라우져를 뛰운다  (2)
        browser.get(url)
        
        # 공지/배너/오버레이 제거 시도(notice_auto_cont 기준)
        try:
            overlay = browser.find_element(By.ID, "notice_auto_cont") 
            try:
                close_btn = overlay.find_element(By.CSS_SELECTOR, ".close, .btn_close, [aria-label='닫기']")
                close_btn.click()
                wait.until(EC.invisibility_of_element_located((By.ID, "notice_auto_cont")))
            except Exception:
                browser.execute_script("arguments[0].style.display='none';", overlay)
                time.sleep(0.1)
        except Exception:
            pass

        # 2.3 url이 loading 될 동안 잠시 기다린다 (3)
        time.sleep(5)

        # 2.4 가상 웹브라우져에서 년도, 조회버튼을 xpath를 이용하여 클릭한다.
        
        # [참고사항]
        # 1) 년도버튼     xpath 주소  : //*[@id="year"]
        # 2) 년도선택버튼 xpath 주소   : /html/body/div[2]/div[3]/div[2]/div/form/select/option[1]
        # 3) 조회버튼     xPath 주소  : //*[@id="txt"]/div[2]/p/span[1]/a/img
        try: # 마우스를 화면 좌상단 빈 곳으로 이동해 hover 해제 
            ActionChains(browser).move_by_offset(-10000, -10000).perform() 
            time.sleep(0.1) 
        except Exception: 
            pass


        
        browser.find_element(By.XPATH,'//*[@id="year"]').click() # 년도 클릭
        #browser.implicitly_wait(1)
        browser.find_element(By.XPATH,'/html/body/div[2]/div[3]/div[2]/div/form/select/option['+str(y_i)+']').click() # 년도 클릭

        browser.find_element(By.XPATH,'//*[@id="txt"]/div[2]/p/span[1]/a/img').click() #  조희 클릭

        # 브라우져가 작동할 시간을 강제적으로 5초정도 준다.
        time.sleep(5)

        html = browser.page_source
        data = pd.read_html(html, header=0)
        #print("data_1",data[1])
        #print("data_2",data[2])
        #print("data_3",data[3])
        #print("data_4",data[4])

        # 추출한 대이터를 정렬한다. 
        data_month_price_raw = data[3]
        data_mprice = data_month_price_raw
        
        # 3. 스크래핑한데이터를 정리(Data Mining)한다.
        #print(data_mprice)
        data_mprice.columns = ["년/월","계약전력(kW)","요금적용전력(kW)","사용전력량(kWh)","day","pf","pf1","전기요금(원)","others"]               # 컬럼명을 1~12월로 변경
        data_mprice.index = [1,2,3,4,5,6,7,8,9,10,11,12]  # 인덱스를 1~12월로 변경
        data_mprice = data_mprice.replace('-','NaN')                       # - 을 NaN으로 변경
        data_mprice = data_mprice.apply(pd.to_numeric,errors='ignore')     # 숫자가 아닌 부분은 무시
        data_mprice = data_mprice.fillna(0)                                # NaN 부분을 0으로 변경
        data_mprice = data_mprice.drop(['day', 'pf', 'pf1','others'], axis=1)
        data_mprice['unit_price'] = (data_mprice['전기요금(원)'] / data_mprice['사용전력량(kWh)']).round(1)
        data_mprice['년/월'] = data_mprice['년/월'].apply(lambda x: x[5:] if '년' in x and len(x) > 5 else x)
        total_usage = data_mprice['사용전력량(kWh)'].sum()
        total_bill = data_mprice['전기요금(원)'].sum()
        average_unit_price = data_mprice['unit_price'].mean().round(1)

        new_row = {
            '년/월': '총계',
            '계약전력(kW)': np.nan,        # NaN으로 표시 (총계에 의미가 없는 값)
            '요금적용전력(kW)': np.nan,     # NaN으로 표시
            '사용전력량(kWh)': total_usage,
            '전기요금(원)': total_bill,
            'unit_price': average_unit_price
        }

        data_mprice = pd.concat([data_mprice, pd.DataFrame([new_row])], ignore_index=True)

        # =================================
        # 월별 단가 1월부터 12월까지 데이터 소팅
        # =================================
        # 1. '년/월' 컬럼의 월 표현을 두 자리 숫자로 통일합니다. (예: '1월' -> '01월')
        def format_month(month_str):
            # "년/월" 칼럼의 순서를 만드는 함수
            if month_str == '총계':
                return month_str
            # '1월', '2월' 등 한 자리 월에 '0'을 붙여 '01월', '02월'로 만듭니다.
            # '10월', '11월', '12월' 등은 이미 두 자리이므로 그대로 둡니다.
            return f"{int(month_str[:-1]):02d}월"

        data_mprice['년/월'] = data_mprice['년/월'].apply(format_month)

        # 2. 월 순서를 정의하는 리스트를 생성합니다. 
        month_order = [f'{i:02d}월' for i in range(1, 13)] + ['총계']

        # 3. DataFrame(data_mprice)의 '년/월' 컬럼을 Categorical 타입으로 변환하면서 정의된 순서(month_order)를 적용합니다.
        data_mprice['년/월'] = pd.Categorical(data_mprice['년/월'], categories=month_order, ordered=True)

        # 4. '년/월' 컬럼을 기준으로 DataFrame(data_mprice)을 정렬합니다.
        data_mprice = data_mprice.sort_values('년/월').reset_index(drop=True) # 정렬 후 인덱스 초기화

        #print(data_mprice)
        # 4. 최대수요량을 저장한다.
        file_name = "./data/"+ user_data_df["User ID"].iloc[0]+"_"+str(search_year)+"_monthly_price.csv"
        data_mprice.to_csv(file_name,encoding='euc-kr')

        # 5. 브라우져 종료
        browser.quit()


if __name__ == '__main__':
    main()
    #app = QApplication([])
    #window = MatplotlibWidget()
    #window.show()
    #app.exec_()
